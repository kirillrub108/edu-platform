"""Unit tests for the script-text extractors in app.routers.uploads.

These are the pure helpers behind POST /api/v1/uploads/script: whatever they
return becomes the lecture narration, so a silently empty or mangled result is
expensive. DOCX zip-bomb guards live in test_docx_extraction.py.
"""

from __future__ import annotations

import io
import subprocess
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest

from app.routers import uploads as uploads_mod
from app.routers.uploads import (
    _decode_text,
    _extract_docx_text,
    _extract_html_text,
    _extract_odt_text,
    _extract_pdf_text,
    _extract_rtf_text,
    _extract_script_text,
    _extract_via_libreoffice,
)

pytestmark = pytest.mark.unit


def _pdf_bytes(pages: list[str]) -> bytes:
    """Build a minimal one-font PDF. pypdf refuses a file without a valid xref
    table, so offsets are tracked and emitted explicitly."""
    objs: dict[int, bytes] = {
        1: b"<</Type/Catalog/Pages 2 0 R>>",
        3: b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    }
    kids = " ".join(f"{4 + 2 * i} 0 R" for i in range(len(pages)))
    objs[2] = f"<</Type/Pages/Kids[{kids}]/Count {len(pages)}>>".encode()
    for i, text in enumerate(pages):
        page_id, content_id = 4 + 2 * i, 5 + 2 * i
        objs[page_id] = (
            f"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents {content_id} 0 R"
            f"/Resources<</Font<</F1 3 0 R>>>>>>"
        ).encode()
        stream = f"BT /F1 24 Tf 72 700 Td ({text}) Tj ET".encode() if text else b""
        objs[content_id] = f"<</Length {len(stream)}>>stream\n".encode() + stream + b"\nendstream"

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += f"{num} 0 obj".encode() + objs[num] + b"endobj\n"

    xref_at = len(out)
    total = max(objs) + 1
    out += f"xref\n0 {total}\n".encode() + b"0000000000 65535 f \n"
    for num in range(1, total):
        out += f"{offsets[num]:010d} 00000 n \n".encode()
    out += f"trailer<</Root 1 0 R/Size {total}>>\nstartxref\n{xref_at}\n%%EOF\n".encode()
    return bytes(out)


def _docx_bytes_with_table() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Вводный абзац")
    doc.add_paragraph("   ")  # whitespace-only paragraphs are dropped
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Скорость"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _odt_bytes() -> bytes:
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P

    doc = OpenDocumentText()
    doc.text.addElement(P(text="Первый абзац"))
    doc.text.addElement(P(text="   "))
    doc.text.addElement(H(outlinelevel=1, text="Заголовок"))
    buf = io.BytesIO()
    doc.write(buf)
    return buf.getvalue()


# ── _decode_text ──────────────────────────────────────────────────────────────


def test_decode_text_reads_utf8_with_bom() -> None:
    assert _decode_text("Лекция".encode("utf-8-sig")) == "Лекция"


def test_decode_text_falls_back_to_cp1251() -> None:
    """Windows-authored .txt scripts are the common case for teachers."""
    assert _decode_text("Лекция по физике".encode("cp1251")) == "Лекция по физике"


def test_decode_text_never_raises_on_undecodable_bytes() -> None:
    # 0x98 is undefined in cp1251 and invalid UTF-8 — the last resort drops it.
    assert _decode_text(b"plain\x98text") == "plaintext"


# ── per-format extractors ─────────────────────────────────────────────────────


def test_extract_pdf_joins_pages_and_skips_empty_ones() -> None:
    text = _extract_pdf_text(_pdf_bytes(["Hello lecture", "", "Second page"]))
    assert text == "Hello lecture\n\nSecond page"


def test_extract_docx_includes_table_rows() -> None:
    text = _extract_docx_text(_docx_bytes_with_table())
    assert "Вводный абзац" in text
    assert "Параметр | Значение" in text
    # A row whose second cell is empty keeps only the filled cell.
    assert "Скорость" in text
    assert "|  |" not in text


def test_extract_rtf_returns_plain_text() -> None:
    rtf = rb"{\rtf1\ansi\ansicpg1251 Hello \b bold\b0  text\par}"
    assert "Hello" in _extract_rtf_text(rtf)
    assert "bold" in _extract_rtf_text(rtf)
    assert "rtf1" not in _extract_rtf_text(rtf)


def test_extract_odt_returns_paragraphs_and_headings() -> None:
    text = _extract_odt_text(_odt_bytes())
    assert "Первый абзац" in text
    assert "Заголовок" in text
    assert text.count("\n\n") == 1  # the blank paragraph was dropped


def test_extract_html_skips_script_style_and_head() -> None:
    html = (
        "<html><head><title>Заголовок вкладки</title>"
        "<style>p{color:red}</style></head>"
        "<body><script>var x = 1;</script>"
        "<h1>Тема</h1><p>Первый абзац</p><div>Второй абзац</div>"
        "<ul><li>Пункт</li></ul></body></html>"
    ).encode()

    text = _extract_html_text(html)

    assert text.splitlines() == ["Тема", "Первый абзац", "Второй абзац", "Пункт"]
    assert "Заголовок вкладки" not in text
    assert "var x" not in text


def test_extract_html_decodes_entities() -> None:
    assert _extract_html_text(b"<p>&laquo;&nbsp;&raquo;&amp;</p>").strip().endswith("&")


# ── _extract_via_libreoffice ──────────────────────────────────────────────────


def _fake_subprocess(behaviour: Any, *, write_output: bool = True) -> SimpleNamespace:
    """Stand-in for the `subprocess` module inside uploads (the except clause
    resolves its exception classes through the same namespace)."""
    calls: list[list[str]] = []

    def _run(cmd: list[str], **_kwargs: Any) -> SimpleNamespace:
        calls.append(cmd)
        if isinstance(behaviour, Exception):
            raise behaviour
        if write_output:
            outdir = cmd[cmd.index("--outdir") + 1]
            Path(outdir, "input.txt").write_bytes("Текст из .doc".encode("cp1251"))
        return SimpleNamespace(returncode=0)

    return SimpleNamespace(
        run=_run,
        calls=calls,
        CalledProcessError=subprocess.CalledProcessError,
        TimeoutExpired=subprocess.TimeoutExpired,
    )


def test_libreoffice_converts_and_decodes_output(monkeypatch: pytest.MonkeyPatch) -> None:
    fake = _fake_subprocess(None)
    monkeypatch.setattr(uploads_mod, "subprocess", fake)

    text = _extract_via_libreoffice(b"\xd0\xcf\x11\xe0legacy doc", ".doc")

    assert text == "Текст из .doc"
    cmd = fake.calls[0]
    assert cmd[0] == "libreoffice"
    assert "--headless" in cmd
    assert cmd[cmd.index("--convert-to") + 1].startswith("txt:")


@pytest.mark.parametrize(
    "exc",
    [
        subprocess.CalledProcessError(1, ["libreoffice"]),
        subprocess.TimeoutExpired(["libreoffice"], 60),
    ],
)
def test_libreoffice_failure_becomes_runtime_error(
    monkeypatch: pytest.MonkeyPatch, exc: Exception
) -> None:
    monkeypatch.setattr(uploads_mod, "subprocess", _fake_subprocess(exc))

    with pytest.raises(RuntimeError, match="LibreOffice conversion failed"):
        _extract_via_libreoffice(b"legacy", ".doc")


def test_libreoffice_silent_no_output_becomes_runtime_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Exit code 0 with no .txt produced is LibreOffice's usual failure mode."""
    monkeypatch.setattr(uploads_mod, "subprocess", _fake_subprocess(None, write_output=False))

    with pytest.raises(RuntimeError, match="produced no output"):
        _extract_via_libreoffice(b"legacy", ".doc")


# ── _extract_script_text dispatch ─────────────────────────────────────────────


@pytest.mark.parametrize("name", ["script.txt", "SCRIPT.MD", "notes.markdown"])
def test_dispatch_plain_text_variants(name: str) -> None:
    assert _extract_script_text(name, "Текст".encode("cp1251")) == "Текст"


def test_dispatch_pdf() -> None:
    assert _extract_script_text("lecture.pdf", _pdf_bytes(["Hello lecture"])) == "Hello lecture"


def test_dispatch_docx() -> None:
    assert "Вводный абзац" in _extract_script_text("lecture.docx", _docx_bytes_with_table())


def test_dispatch_rtf() -> None:
    out = _extract_script_text("lecture.rtf", rb"{\rtf1\ansi Hello\par}")
    assert "Hello" in out


def test_dispatch_odt() -> None:
    assert "Первый абзац" in _extract_script_text("lecture.odt", _odt_bytes())


@pytest.mark.parametrize("name", ["page.html", "page.HTM"])
def test_dispatch_html_variants(name: str) -> None:
    assert _extract_script_text(name, "<p>Абзац</p>".encode()) == "Абзац"


def test_dispatch_legacy_doc_goes_through_libreoffice(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(uploads_mod, "subprocess", _fake_subprocess(None))

    assert _extract_script_text("lecture.doc", b"legacy") == "Текст из .doc"


@pytest.mark.parametrize("name", ["archive.zip", "slides.pptx", "noextension"])
def test_dispatch_rejects_unsupported_extension(name: str) -> None:
    with pytest.raises(ValueError, match="Unsupported extension"):
        _extract_script_text(name, b"data")
