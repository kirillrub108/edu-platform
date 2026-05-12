import io
import os
import subprocess
import tempfile
from html.parser import HTMLParser
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.dependencies import require_teacher
from app.models.lesson import Lesson
from app.models.user import User
from app.services.storage_service import storage_service

router = APIRouter(prefix="/api/v1/uploads", tags=["uploads"])

ALLOWED_PPTX = {".pptx", ".ppt", ".pdf"}
ALLOWED_VIDEO = {".mp4", ".webm", ".mov"}
ALLOWED_SCRIPT = {
    ".txt", ".md", ".markdown",
    ".pdf",
    ".docx", ".doc",
    ".rtf",
    ".odt",
    ".html", ".htm",
}

MAX_SCRIPT_BYTES = 10 * 1024 * 1024  # 10 MB


def _ext_ok(filename: str | None, allowed: set[str]) -> bool:
    if not filename:
        return False
    return os.path.splitext(filename)[1].lower() in allowed


def _decode_text(content: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "windows-1251"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF — text only, no images."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx_text(content: bytes) -> str:
    """Extract paragraphs and table cells from a DOCX. Images/embedded objects are skipped."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_rtf_text(content: bytes) -> str:
    from striprtf.striprtf import rtf_to_text

    return rtf_to_text(_decode_text(content))


def _extract_odt_text(content: bytes) -> str:
    """Extract text paragraphs from an ODT (OpenDocument Text)."""
    from odf import teletype, text as odf_text
    from odf.opendocument import load

    doc = load(io.BytesIO(content))
    parts: list[str] = []
    for elem in doc.getElementsByType(odf_text.P) + doc.getElementsByType(odf_text.H):
        text = teletype.extractText(elem).strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


class _HTMLTextExtractor(HTMLParser):
    _SKIP = {"script", "style", "head", "noscript"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._buf: list[str] = []
        self._depth_skip = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._SKIP:
            self._depth_skip += 1
        elif tag in ("p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._buf.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._depth_skip > 0:
            self._depth_skip -= 1
        elif tag in ("p", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._buf.append("\n")

    def handle_data(self, data: str) -> None:
        if self._depth_skip == 0:
            self._buf.append(data)

    def get_text(self) -> str:
        raw = "".join(self._buf)
        lines = [line.strip() for line in raw.splitlines()]
        return "\n".join(line for line in lines if line)


def _extract_html_text(content: bytes) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(_decode_text(content))
    return parser.get_text()


def _extract_via_libreoffice(content: bytes, suffix: str) -> str:
    """Fallback for legacy .doc — convert to TXT via headless LibreOffice."""
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, f"input{suffix}")
        with open(src, "wb") as f:
            f.write(content)
        try:
            subprocess.run(
                [
                    "libreoffice", "--headless",
                    f"-env:UserInstallation=file://{tmp}/lo_profile",
                    "--convert-to", "txt:Text (encoded):UTF8",
                    "--outdir", tmp, src,
                ],
                check=True, capture_output=True, timeout=60,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise RuntimeError(f"LibreOffice conversion failed: {exc}") from exc

        txt_path = os.path.join(tmp, "input.txt")
        if not os.path.exists(txt_path):
            raise RuntimeError("LibreOffice produced no output")
        with open(txt_path, "rb") as f:
            return _decode_text(f.read())


def _extract_script_text(filename: str, content: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md", ".markdown"):
        return _decode_text(content)
    if ext == ".pdf":
        return _extract_pdf_text(content)
    if ext == ".docx":
        return _extract_docx_text(content)
    if ext == ".rtf":
        return _extract_rtf_text(content)
    if ext == ".odt":
        return _extract_odt_text(content)
    if ext in (".html", ".htm"):
        return _extract_html_text(content)
    if ext == ".doc":
        return _extract_via_libreoffice(content, ext)
    raise ValueError(f"Unsupported extension: {ext}")


@router.post("/pptx")
async def upload_pptx(
    file: UploadFile,
    lesson_id: UUID | None = None,
    user: User = Depends(require_teacher),
    db: AsyncSession = Depends(get_db),
):
    if not _ext_ok(file.filename, ALLOWED_PPTX):
        raise HTTPException(status_code=400, detail="Only PPTX/PPT/PDF allowed")

    relative = await storage_service.save_upload(file, "pptx")

    # Optionally attach the uploaded file to a lesson right away
    if lesson_id is not None:
        lesson = await db.get(Lesson, lesson_id)
        if lesson is None:
            raise HTTPException(status_code=404, detail="Lesson not found")
        lesson.pptx_path = relative
        await db.commit()

    return {
        "file_path": relative,
        "file_url": storage_service.get_url(relative, str(user.id)),
    }


@router.post("/script")
async def upload_script(
    file: UploadFile,
    lesson_id: UUID | None = None,
    user: User = Depends(require_teacher),
    db: AsyncSession = Depends(get_db),
):
    """Upload a TXT/MD/PDF file, extract its text and (optionally) save to lesson.script.

    For PDFs only the text layer is extracted — images and figures are skipped.
    """
    if not _ext_ok(file.filename, ALLOWED_SCRIPT):
        raise HTTPException(
            status_code=400,
            detail="Поддерживаются: TXT, MD, PDF, DOCX, DOC, RTF, ODT, HTML",
        )

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Пустой файл")
    if len(content) > MAX_SCRIPT_BYTES:
        raise HTTPException(status_code=400, detail="Файл слишком большой (>10 МБ)")

    try:
        script = _extract_script_text(file.filename or "", content).strip()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Не удалось извлечь текст: {exc}")

    if not script:
        raise HTTPException(
            status_code=400,
            detail="В файле не найден текст (для отсканированных PDF используйте TXT/MD)",
        )

    if lesson_id is not None:
        lesson = await db.get(Lesson, lesson_id)
        if lesson is None:
            raise HTTPException(status_code=404, detail="Lesson not found")
        lesson.script = script
        await db.commit()

    return {"script": script, "chars": len(script)}


@router.post("/video")
async def upload_video(
    file: UploadFile,
    user: User = Depends(require_teacher),
):
    if not _ext_ok(file.filename, ALLOWED_VIDEO):
        raise HTTPException(status_code=400, detail="Only MP4/WebM/MOV allowed")
    relative = await storage_service.save_upload(file, "videos")
    return {
        "file_path": relative,
        "file_url": storage_service.get_url(relative, str(user.id)),
    }
